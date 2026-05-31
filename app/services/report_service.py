import os
import zipfile
from datetime import UTC, date, datetime
from io import BytesIO
from typing import Any

from docx import Document
from docxtpl import DocxTemplate, RichText

from app.models import Member, MemberSkill, User

# Constants for checkboxes
CHECKED = RichText('☒', font='Segoe UI Symbol')
UNCHECKED = RichText('☐', font='Segoe UI Symbol')

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "assets", "templates")
MEMBER_PROFILE_TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "member_profile_template.docx")
TEMPLATE_PATH = MEMBER_PROFILE_TEMPLATE_PATH  # Backward-compatible alias
EVALUATION_SHEET_TEMPLATE_PATHS = [
    os.path.join(TEMPLATE_DIR, "member_evaluation_sheet_template.docx"),
    os.path.join(TEMPLATE_DIR, "member_evaluation_sheet_template.dotx"),
    os.path.join(TEMPLATE_DIR, "BM-MTEC-NS-03 - Phiếu đánh giá thành viên.dotx"),
]


CLASSIFICATION_LABELS = {
    "EXCELLENT": "Xuất sắc",
    "GOOD": "Tốt",
    "PASSED": "Đạt",
    "NEEDS_IMPROVEMENT": "Cần cải thiện",
    "FAILED": "Không đạt",
}


def format_date(d: date | None) -> str:
    if not d:
        return ""
    return d.strftime("%d/%m/%Y")


def _format_datetime(value: datetime | None = None) -> str:
    value = value or datetime.now(UTC)
    if value.tzinfo is not None:
        value = value.astimezone(UTC).replace(tzinfo=None)
    return value.strftime("%d/%m/%Y %H:%M")


def _format_number(value: Any) -> str:
    if value is None:
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if number.is_integer():
        return str(int(number))
    return f"{number:.2f}".rstrip("0").rstrip(".")


def _format_percent(value: Any) -> str:
    if value is None:
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if 0 <= number <= 1:
        number *= 100
    return f"{number:.2f}%".replace(".00%", "%")


def _text(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def _classification_label(value: Any) -> str:
    code = _text(value)
    if not code:
        return ""
    return CLASSIFICATION_LABELS.get(code, code)


def _blockers_text(blockers: list[Any]) -> str:
    if not blockers:
        return "Không ghi nhận"
    lines: list[str] = []
    for item in blockers:
        if isinstance(item, dict):
            code = item.get("code") or item.get("blockerCode") or "UNKNOWN"
            cap = item.get("cap")
            source = item.get("source")
            title = item.get("title")
            parts = [str(code)]
            if cap:
                parts.append(f"giới hạn xếp loại: {_classification_label(cap)}")
            if source:
                parts.append(f"nguồn: {source}")
            if title:
                parts.append(str(title))
            lines.append(" - ".join(parts))
        else:
            lines.append(str(item))
    return "\n".join(lines)


def _row_value(row: dict[str, Any], key: str) -> str:
    return _text(row.get(key))


def _render_template(template_path: str, context: dict[str, Any]) -> BytesIO:
    doc = DocxTemplate(template_path)
    doc.render(context)
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def _first_existing_template(paths: list[str]) -> str | None:
    for path in paths:
        if os.path.exists(path):
            return path
    return None


def generate_member_profile_docx(member: Member, skills: list[MemberSkill]) -> BytesIO:
    """
    Generates a DOCX profile for a member using a template.
    """
    if not os.path.exists(MEMBER_PROFILE_TEMPLATE_PATH):
        raise FileNotFoundError(f"Template not found at {MEMBER_PROFILE_TEMPLATE_PATH}")

    doc = DocxTemplate(MEMBER_PROFILE_TEMPLATE_PATH)

    # Prepare context
    context = {
        "ho_ten": (member.name or "").upper(),
        "gioi_tinh": member.gender or "",
        "ngay_sinh": format_date(member.dob),
        "mssv": member.mssv or "",
        "khoa": member.khoa or "",
        "chuyen_nganh": member.chuyen_nganh or "",
        "sdt": member.phone or "",
        "email": member.email or "",
        "link_fb": "",  # Link FB is not in our Member model yet, adding as placeholder
        "muc_tieu": member.goal or "",
        "dinh_huong": member.orientation or "",
        "cam_ket_time": "", # Placeholder
        "vi_tri": member.role_title or "",
    }

    # Handle Ban checkboxes
    context["c_ban_cn"] = CHECKED if member.ban == "Ban Công nghệ" else UNCHECKED
    context["c_ban_tt"] = CHECKED if member.ban == "Ban Truyền thông" else UNCHECKED
    context["c_ban_vh"] = CHECKED if member.ban == "Ban Vận hành" else UNCHECKED
    context["c_ban_cnh"] = CHECKED if member.ban == "Ban Chủ nhiệm" else UNCHECKED
    context["c_ban_khac"] = UNCHECKED # Default to unchecked if not one of the above

    # Skill keywords mapping from MTEC_App.py
    skill_map = {
        'tk': 'Thiết kế', 'qd': 'Quay dựng', 'ct': 'Content',
        'fp': 'Fanpage', 'ca': 'Chụp ảnh', 'lt': 'Lập trình', 'mc': 'MC',
        'gt': 'Giao tiếp', 'lvn': 'Làm việc nhóm', 'qltg': 'thời gian',
        'st': 'Sáng tạo', 'gqvd': 'vấn đề', 'thvp': 'văn phòng'
    }

    # Initialize all skill fields to UNCHECKED
    for key in skill_map.keys():
        context[f"{key}_cb"] = UNCHECKED
        context[f"{key}_tb"] = UNCHECKED
        context[f"{key}_tot"] = UNCHECKED
        context[f"{key}"] = UNCHECKED

    # Fill in skills from database
    for skill in skills:
        # Try to find which key this skill belongs to
        found_key = None
        for key, keyword in skill_map.items():
            if keyword.lower() in (skill.name or "").lower():
                found_key = key
                break

        if found_key:
            context[found_key] = CHECKED
            level = (skill.level or "").lower()
            if "cơ bản" in level or "co ban" in level:
                context[f"{found_key}_cb"] = CHECKED
            elif "trung bình" in level or "trung binh" in level:
                context[f"{found_key}_tb"] = CHECKED
            elif "tốt" in level or "tot" in level:
                context[f"{found_key}_tot"] = CHECKED

    # Render document
    doc.render(context)

    # Save to buffer
    target_stream = BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)

    return target_stream


def _evaluation_sheet_context(report: dict[str, Any], actor: User | None = None) -> dict[str, Any]:
    member = report.get("member") or {}
    scores = report.get("scores") or {}
    classification = report.get("classification") or {}
    generated_at = report.get("generatedAt")
    if not isinstance(generated_at, datetime):
        generated_at = datetime.now(UTC)

    final_classification = classification.get("final")
    preliminary_classification = classification.get("preliminary")

    breakdown_rows = [
        {
            "criterion_code": _row_value(item, "criterionCode"),
            "component": _row_value(item, "component"),
            "unit_code": _row_value(item, "unitCode"),
            "raw_score": _format_number(item.get("rawScore")),
            "final_score": _format_number(item.get("finalScore")),
            "max_score": _format_number(item.get("maxScoreSnapshot")),
            "evidence_count": _format_number(item.get("evidenceCount")),
            "calculation_note": _row_value(item, "calculationNote"),
        }
        for item in report.get("breakdowns", [])
    ]

    evidence_rows = [
        {
            "title": _row_value(item, "title"),
            "type": _row_value(item, "evidenceType"),
            "status": _row_value(item, "status"),
            "url": _row_value(item, "url"),
            "description": _row_value(item, "description"),
            "captured_at": _format_datetime(item.get("capturedAt")) if isinstance(item.get("capturedAt"), datetime) else _row_value(item, "capturedAt"),
        }
        for item in report.get("evidence", [])
    ]

    appeal_rows = [
        {
            "criterion_code": _row_value(item, "criterionCode"),
            "appeal_type": _row_value(item, "appealType"),
            "status": _row_value(item, "status"),
            "requested_score": _format_number(item.get("requestedScore")),
            "content": _row_value(item, "content"),
            "resolution_note": _row_value(item, "resolutionNote"),
        }
        for item in report.get("appeals", [])
    ]

    discipline_case_rows = [
        {
            "case_code": _row_value(item, "caseCode"),
            "case_type": _row_value(item, "caseType"),
            "severity": _row_value(item, "severity"),
            "status": _row_value(item, "status"),
            "title": _row_value(item, "title"),
            "blocker_code": _row_value(item, "blockerCode"),
            "point_impact": _format_number(item.get("pointImpact")),
        }
        for item in report.get("disciplineCases", [])
    ]

    context: dict[str, Any] = {
        "cycle_id": report.get("cycleId") or "",
        "cycle_code": report.get("cycleCode") or "",
        "cycle_name": report.get("cycleName") or "",
        "cycle_status": report.get("status") or "",
        "report_version": report.get("reportVersion") or "",
        "generated_at": _format_datetime(generated_at),
        "generated_day": generated_at.strftime("%d"),
        "generated_month": generated_at.strftime("%m"),
        "generated_year": generated_at.strftime("%Y"),
        "exported_by": getattr(actor, "full_name", None) or getattr(actor, "username", None) or "",
        "exported_by_username": getattr(actor, "username", None) or "",
        "member_id": member.get("id") or "",
        "mssv": member.get("mssv") or "",
        "ho_ten": member.get("name") or "",
        "ban": member.get("ban") or "",
        "unit_code": member.get("unitCode") or "",
        "role_title": member.get("roleTitle") or "",
        "member_status": member.get("status") or "",
        "component_i_score": _format_number(scores.get("componentI")),
        "component_ii_score": _format_number(scores.get("componentII")),
        "component_iii_a_score": _format_number(scores.get("componentIIIa")),
        "component_iii_b_score": _format_number(scores.get("componentIIIb")),
        "total_score": _format_number(scores.get("total")),
        "attendance_rate": _format_percent(scores.get("attendanceRate")),
        "preliminary_classification": preliminary_classification or "",
        "preliminary_classification_label": _classification_label(preliminary_classification),
        "final_classification": final_classification or "",
        "final_classification_label": _classification_label(final_classification),
        "blockers_text": _blockers_text(report.get("blockers", [])),
        "breakdown_rows": breakdown_rows,
        "evidence_rows": evidence_rows,
        "appeal_rows": appeal_rows,
        "discipline_case_rows": discipline_case_rows,
        "breakdown_count": len(breakdown_rows),
        "evidence_count": len(evidence_rows),
        "appeal_count": len(appeal_rows),
        "discipline_case_count": len(discipline_case_rows),
        "c_xuat_sac": CHECKED if final_classification == "EXCELLENT" else UNCHECKED,
        "c_tot": CHECKED if final_classification == "GOOD" else UNCHECKED,
        "c_dat": CHECKED if final_classification == "PASSED" else UNCHECKED,
        "c_can_cai_thien": CHECKED if final_classification == "NEEDS_IMPROVEMENT" else UNCHECKED,
        "c_khong_dat": CHECKED if final_classification == "FAILED" else UNCHECKED,
    }
    return context


def _generate_member_evaluation_sheet_fallback_docx(context: dict[str, Any]) -> BytesIO:
    doc = Document()
    doc.add_heading("PHIẾU ĐÁNH GIÁ THÀNH VIÊN", 0)

    doc.add_paragraph(f"Mã chu kỳ: {context['cycle_code']} - {context['cycle_name']}")
    doc.add_paragraph(f"Trạng thái chu kỳ: {context['cycle_status']}")
    doc.add_paragraph(f"Thời điểm xuất: {context['generated_at']}")
    if context.get("exported_by"):
        doc.add_paragraph(f"Người xuất: {context['exported_by']} ({context['exported_by_username']})")

    doc.add_heading("1. Thông tin thành viên", level=1)
    member_table = doc.add_table(rows=0, cols=2)
    for label, key in (
        ("Họ và tên", "ho_ten"),
        ("MSSV", "mssv"),
        ("Ban/Đơn vị", "unit_code"),
        ("Chức vụ/Vai trò", "role_title"),
        ("Trạng thái thành viên", "member_status"),
    ):
        row = member_table.add_row().cells
        row[0].text = label
        row[1].text = _text(context.get(key))

    doc.add_heading("2. Kết quả tổng hợp", level=1)
    score_table = doc.add_table(rows=1, cols=2)
    score_table.rows[0].cells[0].text = "Cấu phần"
    score_table.rows[0].cells[1].text = "Điểm/Kết quả"
    for label, key in (
        ("I. Kỷ luật & Chuyên cần", "component_i_score"),
        ("II. Thái độ & Ý thức tổ chức", "component_ii_score"),
        ("III-A. Hiệu suất chuyên môn chung", "component_iii_a_score"),
        ("III-B. Hiệu suất chuyên môn theo Ban", "component_iii_b_score"),
        ("Tổng điểm", "total_score"),
        ("Tỷ lệ chuyên cần", "attendance_rate"),
        ("Xếp loại sơ bộ", "preliminary_classification_label"),
        ("Xếp loại cuối cùng", "final_classification_label"),
    ):
        row = score_table.add_row().cells
        row[0].text = label
        row[1].text = _text(context.get(key))

    doc.add_heading("3. Yếu tố giới hạn/xử lý đặc biệt", level=1)
    doc.add_paragraph(context.get("blockers_text") or "Không ghi nhận")

    doc.add_heading("4. Bảng điểm chi tiết", level=1)
    breakdown_table = doc.add_table(rows=1, cols=8)
    for idx, heading in enumerate(
        ["Mã tiêu chí", "Cấu phần", "Đơn vị", "Điểm thô", "Điểm cuối", "Điểm tối đa", "MC", "Ghi chú"]
    ):
        breakdown_table.rows[0].cells[idx].text = heading
    for item in context.get("breakdown_rows", []):
        row = breakdown_table.add_row().cells
        row[0].text = item["criterion_code"]
        row[1].text = item["component"]
        row[2].text = item["unit_code"]
        row[3].text = item["raw_score"]
        row[4].text = item["final_score"]
        row[5].text = item["max_score"]
        row[6].text = item["evidence_count"]
        row[7].text = item["calculation_note"]

    if context.get("evidence_rows"):
        doc.add_heading("5. Minh chứng", level=1)
        evidence_table = doc.add_table(rows=1, cols=4)
        for idx, heading in enumerate(["Tiêu đề", "Loại", "Trạng thái", "Mô tả/URL"]):
            evidence_table.rows[0].cells[idx].text = heading
        for item in context["evidence_rows"]:
            row = evidence_table.add_row().cells
            row[0].text = item["title"]
            row[1].text = item["type"]
            row[2].text = item["status"]
            row[3].text = item["description"] or item["url"]

    if context.get("appeal_rows"):
        doc.add_heading("6. Phúc khảo/Khiếu nại", level=1)
        appeal_table = doc.add_table(rows=1, cols=5)
        for idx, heading in enumerate(["Tiêu chí", "Loại", "Trạng thái", "Điểm đề nghị", "Nội dung"]):
            appeal_table.rows[0].cells[idx].text = heading
        for item in context["appeal_rows"]:
            row = appeal_table.add_row().cells
            row[0].text = item["criterion_code"]
            row[1].text = item["appeal_type"]
            row[2].text = item["status"]
            row[3].text = item["requested_score"]
            row[4].text = item["content"]

    if context.get("discipline_case_rows"):
        doc.add_heading("7. Hồ sơ kỷ luật liên quan", level=1)
        case_table = doc.add_table(rows=1, cols=6)
        for idx, heading in enumerate(["Mã vụ việc", "Loại", "Mức độ", "Trạng thái", "Ảnh hưởng điểm", "Tiêu đề"]):
            case_table.rows[0].cells[idx].text = heading
        for item in context["discipline_case_rows"]:
            row = case_table.add_row().cells
            row[0].text = item["case_code"]
            row[1].text = item["case_type"]
            row[2].text = item["severity"]
            row[3].text = item["status"]
            row[4].text = item["point_impact"]
            row[5].text = item["title"]

    doc.add_paragraph("")
    doc.add_paragraph("XÁC NHẬN CỦA BAN VẬN HÀNH")
    doc.add_paragraph("(Ký và ghi rõ họ tên)")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


def generate_member_evaluation_sheet_docx(
    report: dict[str, Any], actor: User | None = None
) -> BytesIO:
    """
    Generates a member evaluation sheet DOCX.

    The function prefers a docxtpl template when a BM-MTEC-NS-03 template is present in
    app/assets/templates. If no template is available, it falls back to a structured DOCX
    so the export endpoint remains usable in every environment.
    """
    context = _evaluation_sheet_context(report, actor)
    template_path = _first_existing_template(EVALUATION_SHEET_TEMPLATE_PATHS)
    if template_path:
        return _render_template(template_path, context)
    return _generate_member_evaluation_sheet_fallback_docx(context)


def generate_members_zip(members_with_skills: list[tuple[Member, list[MemberSkill]]]) -> BytesIO:
    """
    Generates a ZIP file containing multiple member profile DOCX files.
    """
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        for member, skills in members_with_skills:
            try:
                docx_buffer = generate_member_profile_docx(member, skills)
                filename = f"HOSO_{member.mssv}_{member.name.replace(' ', '_')}.docx"
                zip_file.writestr(filename, docx_buffer.getvalue())
            except Exception:
                # Log or handle individual file errors if needed
                continue

    zip_buffer.seek(0)
    return zip_buffer
