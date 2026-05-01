from fastapi import APIRouter, Depends, Body
from sqlalchemy.orm import Session
from fastapi.responses import StreamingResponse
import io

from app.core.errors import raise_api_error
from app.core.rate_limit import rate_limiter
from app.core.response import api_response
from app.db import get_db
from app.deps import get_current_user
from app.models import AIGenerationLog, User
from app.schemas import AIGenerateDraftBody, AIGenerateInsightBody, AIProcessContextBody, AIExportDocumentBody
from app.services.ai_provider import AIProviderError, generate_text

router = APIRouter(prefix="/api/ai", tags=["ai"])


def _persist_ai_log(
    db: Session,
    user: User,
    module: str,
    prompt: str,
    response_text: str,
    provider: str,
    status: str = "success",
) -> AIGenerationLog:
    log = AIGenerationLog(
        user_id=user.id,
        module=module,
        prompt=prompt,
        response_text=response_text,
        provider=provider,
        status=status,
    )
    db.add(log)
    db.commit()
    db.refresh(log)
    return log


@router.post(
    "/generate-insight",
    dependencies=[Depends(rate_limiter(max_requests=10, window_seconds=60))],
)
def generate_insight(
    body: AIGenerateInsightBody,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    try:
        text, provider = generate_text(prompt=body.prompt)
        log = _persist_ai_log(
            db=db,
            user=current_user,
            module="dashboard",
            prompt=body.prompt,
            response_text=text,
            provider=provider,
        )
    except AIProviderError as exc:
        _persist_ai_log(
            db=db,
            user=current_user,
            module="dashboard",
            prompt=body.prompt,
            response_text=exc.message,
            provider="gemini",
            status="failed",
        )
        raise_api_error(status_code=exc.status_code, code=exc.code, message=exc.message)

    return api_response(
        data={
            "text": text,
            "logId": log.id,
            "provider": log.provider,
            "status": log.status,
        }
    )


@router.get("/templates")
def list_templates(_: User = Depends(get_current_user)) -> dict:
    # Trả về danh sách các mẫu văn bản hỗ trợ
    templates = [
        {
            "id": "BM-MTEC-HC-03",
            "name": "Biên bản họp CLB",
            "description": "Dùng để ghi lại nội dung các cuộc họp chính thức của CLB.",
        },
        {
            "id": "BM-MTEC-NS-01",
            "name": "Đơn xin nghỉ phép",
            "description": "Dùng cho thành viên xin nghỉ phép có lý do.",
        },
    ]
    return api_response(data=templates)


@router.post("/process-context")
def process_context(
    body: AIProcessContextBody,
    current_user: User = Depends(get_current_user),
) -> dict:
    # Xử lý ngữ cảnh từ file hoặc link. Ở bản dev này, chúng ta giả lập việc trích xuất text.
    # Trong thực tế, sẽ gọi service để parse PDF/Docx hoặc Google Sheets API.
    return api_response(
        data={
            "message": f"Đã xử lý ngữ cảnh từ {body.source} thành công.",
            "extractedLength": len(body.content) * 2,  # Mock length
            "preview": body.content[:100] + "...",
        }
    )


@router.post("/export-document")
def export_document(
    body: AIExportDocumentBody,
    current_user: User = Depends(get_current_user),
):
    # Xuất tài liệu ra DOCX. Ở đây chúng ta giả lập việc tạo file.
    # Sử dụng logic tương tự như report_service.py nếu có template thực tế.
    from docx import Document

    doc = Document()
    doc.add_heading(f"Tài liệu: {body.templateId}", 0)
    doc.add_paragraph(body.content)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"AI_Export_{body.templateId}_{current_user.username}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@router.post(
    "/generate-draft",
    dependencies=[Depends(rate_limiter(max_requests=10, window_seconds=60))],
)
def generate_draft(
    body: AIGenerateDraftBody,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> dict:
    try:
        text, provider = generate_text(prompt=body.prompt, context=body.context)
        log = _persist_ai_log(
            db=db,
            user=current_user,
            module="generator",
            prompt=body.prompt,
            response_text=text,
            provider=provider,
        )
    except AIProviderError as exc:
        _persist_ai_log(
            db=db,
            user=current_user,
            module="generator",
            prompt=body.prompt,
            response_text=exc.message,
            provider="gemini",
            status="failed",
        )
        raise_api_error(status_code=exc.status_code, code=exc.code, message=exc.message)

    return api_response(
        data={
            "text": text,
            "logId": log.id,
            "provider": log.provider,
            "status": log.status,
        }
    )
